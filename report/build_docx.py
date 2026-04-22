#!/usr/bin/env python3
"""Build report.docx from report.md using python-docx with precise formatting."""

import copy
import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPORT_DIR = Path(__file__).parent
MD_FILE = REPORT_DIR / "report.md"
FIGURES_DIR = REPORT_DIR / "figures"
TEMPLATE_FILE = REPORT_DIR.parent / "课程设计报告模板.docx"
OUTPUT_FILE = REPORT_DIR / "report.docx"

# scienceplots figures use 12cm, everything else 15cm
SCIENCEPLOTS_FIGURES = {
    "zipf_distribution.png",
    "df_distribution.png",
    "doc_length_distribution.png",
    "top20_terms.png",
}


# ── helpers ──────────────────────────────────────────────────────────────────

def _ensure_rPr(element):
    rPr = element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        element.insert(0, rPr)
    return rPr


def _set_rFonts(rPr, ascii="Times New Roman", east_asia="宋体"):
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rPr.insert(0, rf)
    # 移除 theme 引用，避免主题字体覆盖显式设置
    for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        if rf.get(qn("w:" + attr)) is not None:
            del rf.attrib[qn("w:" + attr)]
    if ascii:
        rf.set(qn("w:ascii"), ascii)
        rf.set(qn("w:hAnsi"), ascii)
        rf.set(qn("w:cs"), ascii)
    if east_asia:
        rf.set(qn("w:eastAsia"), east_asia)


def _set_sz(rPr, half_points):
    """Set font size in half-points (e.g. 24 = 12pt)."""
    for tag in ("w:sz", "w:szCs"):
        el = rPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rPr.append(el)
        el.set(qn("w:val"), str(half_points))


def _set_bold(rPr, val=True):
    b = rPr.find(qn("w:b"))
    if val:
        if b is None:
            b = OxmlElement("w:b")
            rPr.append(b)
    else:
        if b is not None:
            rPr.remove(b)


def _set_italic(rPr, val=True):
    i = rPr.find(qn("w:i"))
    if val:
        if i is None:
            i = OxmlElement("w:i")
            rPr.append(i)
    else:
        if i is not None:
            rPr.remove(i)


def _ensure_pPr(para):
    pPr = para._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        para._element.insert(0, pPr)
    return pPr


def _set_spacing(pPr, before=None, after=None, line=None):
    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        pPr.append(sp)
    if before is not None:
        sp.set(qn("w:before"), str(before))
    if after is not None:
        sp.set(qn("w:after"), str(after))
    if line is not None:
        sp.set(qn("w:line"), str(line))
        sp.set(qn("w:lineRule"), "auto")


def _set_jc(pPr, val):
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        pPr.append(jc)
    jc.set(qn("w:val"), val)


def _set_indent(pPr, first_line_chars=None, first_line=None):
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    if first_line_chars is not None:
        ind.set(qn("w:firstLineChars"), str(first_line_chars))
    if first_line is not None:
        ind.set(qn("w:firstLine"), str(first_line))


def _clear_indent(pPr):
    for ind in pPr.findall(qn("w:ind")):
        pPr.remove(ind)
    # 显式归零以覆盖父样式（Normal / BodyText）可能带来的首行缩进
    ind = OxmlElement("w:ind")
    ind.set(qn("w:firstLineChars"), "0")
    ind.set(qn("w:firstLine"), "0")
    pPr.append(ind)


def make_run(text, ascii="Times New Roman", east_asia="宋体", sz=24,
             bold=False, italic=False):
    """Create a w:r element with formatting."""
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    r.insert(0, rPr)
    _set_rFonts(rPr, ascii=ascii, east_asia=east_asia)
    _set_sz(rPr, sz)
    if bold:
        _set_bold(rPr)
    if italic:
        _set_italic(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


# ── paragraph builders ───────────────────────────────────────────────────────

def _add_inline_runs(para, text, default_sz=24, default_ascii="Times New Roman",
                     default_east="宋体"):
    """Parse inline markdown (**bold**, `code`, $math$) into runs."""
    pattern = re.compile(r'(\*\*(.+?)\*\*|`([^`]+)`|\$([^$]+)\$)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            para._element.append(make_run(text[pos:m.start()], sz=default_sz,
                                          ascii=default_ascii, east_asia=default_east))
        if m.group(2):  # bold
            para._element.append(make_run(m.group(2), sz=default_sz, bold=True,
                                          ascii=default_ascii, east_asia=default_east))
        elif m.group(3):  # inline code
            para._element.append(make_run(m.group(3), sz=default_sz,
                                          ascii="Courier New", east_asia=default_east))
        elif m.group(4):  # inline math
            para._element.append(make_run(m.group(4), sz=default_sz, italic=True,
                                          ascii=default_ascii, east_asia=default_east))
        pos = m.end()
    if pos < len(text):
        para._element.append(make_run(text[pos:], sz=default_sz,
                                      ascii=default_ascii, east_asia=default_east))


def add_body_para(doc, text):
    """Normal body paragraph: 12pt, first-line indent 2 chars."""
    p = doc.add_paragraph()
    pPr = _ensure_pPr(p)
    _set_indent(pPr, first_line_chars=200, first_line=480)
    _add_inline_runs(p, text)
    return p


def add_unindented_para(doc, text):
    """Paragraph without indent (list items etc.)."""
    p = doc.add_paragraph()
    pPr = _ensure_pPr(p)
    _clear_indent(pPr)
    _add_inline_runs(p, text)
    return p


def add_heading2(doc, text):
    """H2: sz=32 (16pt), bold, 宋体+TNR, before=260 after=260 line=416, no indent."""
    p = doc.add_paragraph()
    pPr = _ensure_pPr(p)
    _set_spacing(pPr, before=260, after=260, line=416)
    _clear_indent(pPr)
    p._element.append(make_run(text, sz=32, bold=True))
    return p


def add_heading3(doc, text):
    """H3: sz=32 (16pt), bold, before=260 after=260 line=416, no indent."""
    p = doc.add_paragraph()
    pPr = _ensure_pPr(p)
    _set_spacing(pPr, before=260, after=260, line=416)
    _clear_indent(pPr)
    p._element.append(make_run(text, sz=32, bold=True))
    return p


def add_heading4(doc, text):
    """H4: sz=28 (14pt), bold, 宋体+TNR, before=280 after=290 line=376, no indent."""
    p = doc.add_paragraph()
    pPr = _ensure_pPr(p)
    _set_spacing(pPr, before=280, after=290, line=376)
    _clear_indent(pPr)
    p._element.append(make_run(text, sz=28, bold=True))
    return p


def add_abstract_title(doc, text):
    """Abstract title: sz=28 (14pt), bold, 黑体, centered."""
    p = doc.add_paragraph()
    pPr = _ensure_pPr(p)
    _set_jc(pPr, "center")
    _clear_indent(pPr)
    p._element.append(make_run(text, sz=28, bold=True, east_asia="黑体"))
    return p


def add_figure(doc, img_path, caption):
    """Insert image centered + caption below."""
    filename = os.path.basename(img_path)
    width = Cm(12) if filename in SCIENCEPLOTS_FIGURES else Cm(15)

    # Image paragraph
    p = doc.add_paragraph()
    pPr = _ensure_pPr(p)
    _set_jc(pPr, "center")
    _clear_indent(pPr)
    run = p.add_run()
    run.add_picture(str(img_path), width=width)

    # Caption paragraph: sz=21 (10.5pt), centered
    cap = doc.add_paragraph()
    cpPr = _ensure_pPr(cap)
    _set_jc(cpPr, "center")
    _clear_indent(cpPr)
    cap._element.append(make_run(caption, sz=21))
    return p


def add_code_block(doc, lines):
    """Code block: Courier New 9pt (sz=18), no indent, tight spacing."""
    for line in lines:
        p = doc.add_paragraph()
        pPr = _ensure_pPr(p)
        _clear_indent(pPr)
        _set_spacing(pPr, before=0, after=0, line=260)
        p._element.append(make_run(line if line else " ", sz=18,
                                    ascii="Courier New", east_asia="宋体"))


def add_math_block(doc, text):
    """Display math: italic, centered, 12pt."""
    p = doc.add_paragraph()
    pPr = _ensure_pPr(p)
    _set_jc(pPr, "center")
    _clear_indent(pPr)
    p._element.append(make_run(text, sz=24, italic=True))
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    pPr = _ensure_pPr(p)
    _clear_indent(pPr)
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p._element.append(r)
    return p


# ── page / style setup ───────────────────────────────────────────────────────

def setup_page(doc):
    s = doc.sections[0]
    s.page_width = Twips(11906)
    s.page_height = Twips(16838)
    # 贴合示例报告：上下 2.54cm、左右对称 3.18cm
    s.top_margin = Twips(1440)
    s.bottom_margin = Twips(1440)
    s.left_margin = Twips(1800)
    s.right_margin = Twips(1800)
    s.header_distance = Twips(851)
    s.footer_distance = Twips(850)

    # sectPr: docGrid type=lines linePitch=312、页码起始 1
    sectPr = s._sectPr
    docGrid = sectPr.find(qn("w:docGrid"))
    if docGrid is None:
        docGrid = OxmlElement("w:docGrid")
        sectPr.append(docGrid)
    docGrid.set(qn("w:type"), "lines")
    docGrid.set(qn("w:linePitch"), "312")

    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)
    pgNumType.set(qn("w:start"), "1")


def setup_header_footer(doc):
    """按示例报告格式添加页眉/页脚。"""
    section = doc.sections[0]

    # —— 页眉：华北理工大学 理学院，居中，底边框 ——
    header = section.header
    # 清掉默认段落
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)
    hp = header.add_paragraph()
    hpPr = _ensure_pPr(hp)
    _set_jc(hpPr, "center")
    _clear_indent(hpPr)
    # 底边框
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    hpPr.insert(0, pBdr)
    hp._element.append(make_run("华北理工大学 理学院", sz=21))

    # —— 页脚：PAGE 域，居中 ——
    footer = section.footer
    for p in list(footer.paragraphs):
        p._element.getparent().remove(p._element)
    fp = footer.add_paragraph()
    fpPr = _ensure_pPr(fp)
    _set_jc(fpPr, "center")
    _clear_indent(fpPr)
    fr = OxmlElement("w:r")
    frPr = OxmlElement("w:rPr")
    fr.append(frPr)
    _set_rFonts(frPr, ascii="Times New Roman", east_asia="宋体")
    _set_sz(frPr, 21)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE \\* MERGEFORMAT")
    # fldSimple 需要包含结果 run
    inner_r = OxmlElement("w:r")
    inner_rPr = OxmlElement("w:rPr")
    inner_r.append(inner_rPr)
    _set_rFonts(inner_rPr, ascii="Times New Roman", east_asia="宋体")
    _set_sz(inner_rPr, 21)
    inner_t = OxmlElement("w:t")
    inner_t.text = "1"
    inner_r.append(inner_t)
    fld.append(inner_r)
    fp._element.append(fld)


def setup_normal_style(doc):
    """Normal: 12pt/sz24, 宋体+TNR, line=360/auto, jc=both."""
    style_el = doc.styles["Normal"].element
    # rPr
    rPr = _ensure_rPr(style_el)
    _set_rFonts(rPr, ascii="Times New Roman", east_asia="宋体")
    _set_sz(rPr, 24)
    # pPr
    pPr = style_el.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        style_el.append(pPr)
    _set_spacing(pPr, line=360)
    _set_jc(pPr, "both")


# ── markdown parser ──────────────────────────────────────────────────────────

def parse_md(text):
    """Return list of (type, content) tuples."""
    lines = text.split("\n")
    elems = []
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]

        # blank
        if not line.strip():
            i += 1
            continue

        # hr – skip
        if line.strip() == "---":
            i += 1
            continue

        # page break
        if line.strip() == "\\newpage":
            elems.append(("pagebreak", None))
            i += 1
            continue

        # headings (order matters: check #### before ### before ##)
        if line.startswith("#### "):
            elems.append(("h4", line[5:].strip()))
            i += 1
            continue
        if line.startswith("### "):
            elems.append(("h3", line[4:].strip()))
            i += 1
            continue
        if line.startswith("## "):
            title = line[3:].strip()
            if title.startswith("摘") or title == "Abstract":
                elems.append(("abstract_title", title))
            else:
                elems.append(("h2", title))
            i += 1
            continue

        # figure
        fm = re.match(r'^!\[(图\s*\d+\s+.+?)\]\((.+?)\)$', line)
        if fm:
            elems.append(("figure", (fm.group(1), fm.group(2))))
            i += 1
            continue

        # code block
        if line.strip().startswith("```"):
            code = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            if i < n:
                i += 1  # skip closing ```
            elems.append(("code", code))
            continue

        # display math $$...$$
        if line.strip().startswith("$$"):
            content = line.strip()[2:]
            if content.endswith("$$") and len(content) > 2:
                elems.append(("math", content[:-2].strip()))
                i += 1
                continue
            parts = []
            if content.strip():
                parts.append(content.strip())
            i += 1
            while i < n:
                l = lines[i].strip()
                if l.endswith("$$"):
                    rest = l[:-2].strip()
                    if rest:
                        parts.append(rest)
                    i += 1
                    break
                parts.append(l)
                i += 1
            elems.append(("math", " ".join(parts)))
            continue

        # list item
        if re.match(r'^(\d+\.\s+|- )', line):
            elems.append(("list_item", line.strip()))
            i += 1
            continue

        # normal paragraph – collect consecutive non-special lines
        para_lines = [line]
        i += 1
        while i < n:
            nl = lines[i]
            if not nl.strip():
                break
            if (nl.startswith("#") or nl.startswith("!") or nl.startswith("```") or
                nl.strip().startswith("$$") or nl.strip() == "---" or
                nl.strip() == "\\newpage" or re.match(r'^(\d+\.\s+|- )', nl)):
                break
            para_lines.append(nl)
            i += 1
        elems.append(("paragraph", " ".join(l.strip() for l in para_lines)))

    return elems


# ── cover page from template ─────────────────────────────────────────────────

def insert_cover(doc, template_path):
    """Copy first 14 paragraphs from template into doc, fix title + color, then page break."""
    from docx.opc.part import Part
    from docx.opc.package import OpcPackage
    from docx.opc.packuri import PackURI

    tmpl = Document(str(template_path))

    # Collect template image blobs keyed by rId, create NEW image parts with
    # unique names (prefixed "cover_") to avoid collisions with body images.
    rid_map = {}  # old_rid -> new_rid
    for rid, rel in tmpl.part.rels.items():
        if "image" in rel.reltype:
            blob = rel.target_part.blob
            content_type = rel.target_part.content_type
            # Create a new part with a unique name
            old_name = rel.target_part.partname  # e.g. /word/media/image1.png
            ext = os.path.splitext(str(old_name))[1]
            # Find next free image number in doc
            existing = {str(r.target_part.partname) for r2, r in doc.part.rels.items()
                        if "image" in r.reltype} if doc.part.rels else set()
            num = 100  # start at high number to avoid collisions
            while f"/word/media/image{num}{ext}" in existing:
                num += 1
            new_partname = PackURI(f"/word/media/image{num}{ext}")
            new_part = Part(new_partname, content_type, blob, doc.part.package)
            new_rid = doc.part.relate_to(new_part, rel.reltype)
            rid_map[rid] = new_rid

    body = doc.element.body

    # 只复制 12 段（到"设计时间"为止），去掉末尾空段避免封面溢出到下一页
    for idx in range(12):
        src = tmpl.paragraphs[idx]
        elem = copy.deepcopy(src._element)

        # Remap image rIds
        for blip in elem.findall(".//" + qn("a:blip")):
            old_rid = blip.get(qn("r:embed"))
            if old_rid and old_rid in rid_map:
                blip.set(qn("r:embed"), rid_map[old_rid])

        # Paragraph 4: replace title placeholder, fix color
        if idx == 4:
            title_set = False
            for r in elem.findall(qn("w:r")):
                for t in r.findall(qn("w:t")):
                    if t.text and ("此处" in t.text or "仿宋" in t.text or "题目" in t.text):
                        if not title_set:
                            t.text = "基于Cranfield数据集的信息检索系统"
                            title_set = True
                        else:
                            t.text = ""
                    elif t.text and t.text.strip() == "" and title_set:
                        t.text = ""
                rPr = r.find(qn("w:rPr"))
                if rPr is not None:
                    col = rPr.find(qn("w:color"))
                    if col is not None and col.get(qn("w:val")) == "FF0000":
                        col.set(qn("w:val"), "000000")

        body.insert(idx, elem)

    # Page break after cover：附加到封面最后一段末尾，避免插入"只含分页符
    # 的空段落"被当作独立一行而产生空白页。
    last_cover_p = body[11]
    break_r = OxmlElement("w:r")
    break_br = OxmlElement("w:br")
    break_br.set(qn("w:type"), "page")
    break_r.append(break_br)
    last_cover_p.append(break_r)


# ── pandoc 转换 + 后处理 ─────────────────────────────────────────────────────

def run_pandoc(md_file, out_file):
    """Markdown → docx，交由 pandoc 处理公式/图片/caption/目录。"""
    import subprocess
    cmd = [
        "pandoc", str(md_file),
        "--from", "markdown+tex_math_dollars",
        "--to", "docx",
        "--toc", "--toc-depth=3",
        "--metadata", "toc-title=目录",
        "--output", str(out_file),
    ]
    subprocess.run(cmd, check=True, cwd=md_file.parent)


def _find_style_by_id(doc, style_id):
    for s in doc.styles:
        if s.element.get(qn("w:styleId")) == style_id:
            return s
    return None


def _strip_color(rPr):
    for c in rPr.findall(qn("w:color")):
        rPr.remove(c)


def override_heading_styles(doc):
    """覆盖 Heading1-4 的字体/字号/粗细/间距/颜色（去蓝色）。"""
    levels = [
        # (style_id, sz_half_pt, before_twips, after_twips, line_twips)
        # 示例对应值：H2=32/260/260/416, H3=32/260/260/416, H4=28/280/290/376
        ("Heading1", 32, 260, 260, 416),
        ("Heading2", 32, 260, 260, 416),
        ("Heading3", 32, 260, 260, 416),
        ("Heading4", 28, 280, 290, 376),
    ]
    for style_id, sz, before, after, line in levels:
        style = _find_style_by_id(doc, style_id)
        if style is None:
            continue
        el = style.element
        rPr = _ensure_rPr(el)
        _strip_color(rPr)
        _set_rFonts(rPr, ascii="Times New Roman", east_asia="宋体")
        _set_sz(rPr, sz)
        _set_bold(rPr)
        _set_italic(rPr, False)
        # 同步 Char 样式（Heading1Char/Heading2Char…）
        char_style = _find_style_by_id(doc, style_id + "Char")
        if char_style is not None:
            cel = char_style.element
            crPr = _ensure_rPr(cel)
            _strip_color(crPr)
            _set_rFonts(crPr, ascii="Times New Roman", east_asia="宋体")
            _set_sz(crPr, sz)
            _set_bold(crPr)
            _set_italic(crPr, False)
        pPr = el.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            el.insert(0, pPr)
        _set_spacing(pPr, before=before, after=after, line=line)
        _clear_indent(pPr)
        # 与下段同页 + 段中不分页：避免标题悬挂在页末
        for tag in ("w:keepNext", "w:keepLines"):
            if pPr.find(qn(tag)) is None:
                pPr.append(OxmlElement(tag))
    # TOCHeading 也去蓝色 + 居中
    toch = _find_style_by_id(doc, "TOCHeading")
    if toch is not None:
        rPr = _ensure_rPr(toch.element)
        _strip_color(rPr)
        _set_rFonts(rPr, ascii="Times New Roman", east_asia="黑体")
        _set_sz(rPr, 28)
        _set_bold(rPr)
        pPr = toch.element.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            toch.element.insert(0, pPr)
        _set_jc(pPr, "center")
        _clear_indent(pPr)
    # Hyperlink 样式改为黑色无下划线（TOC 用它）；正文 [N] 自带 inline color 覆盖
    hl = _find_style_by_id(doc, "Hyperlink")
    if hl is not None:
        rPr = _ensure_rPr(hl.element)
        _strip_color(rPr)
        # 显式 auto color
        col = OxmlElement("w:color")
        col.set(qn("w:val"), "auto")
        rPr.append(col)
        # 去下划线
        for u in rPr.findall(qn("w:u")):
            rPr.remove(u)

    # Caption / ImageCaption / CaptionedFigure：去 italic、居中、去缩进
    for sid in ("Caption", "ImageCaption", "CaptionedFigure"):
        style = _find_style_by_id(doc, sid)
        if style is None:
            continue
        rPr = _ensure_rPr(style.element)
        _strip_color(rPr)
        _set_italic(rPr, False)
        if sid != "CaptionedFigure":
            _set_rFonts(rPr, ascii="Times New Roman", east_asia="宋体")
            _set_sz(rPr, 21)
        pPr = style.element.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            style.element.insert(0, pPr)
        _set_jc(pPr, "center")
        _clear_indent(pPr)


def override_body_styles(doc):
    """所有正文样式统一：宋体+TNR, sz=24, line=360, 首行缩进 2 字符, 两端对齐。"""
    for sid in ("Normal", "BodyText", "FirstParagraph", "Compact"):
        style = _find_style_by_id(doc, sid)
        if style is None:
            continue
        el = style.element
        rPr = _ensure_rPr(el)
        _set_rFonts(rPr, ascii="Times New Roman", east_asia="宋体")
        _set_sz(rPr, 24)
        pPr = el.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            el.insert(0, pPr)
        _set_spacing(pPr, line=360)
        _set_indent(pPr, first_line_chars=200, first_line=480)
        _set_jc(pPr, "both")


def override_code_style(doc):
    """SourceCode：Courier 紧凑行距、左对齐、不缩进、不换行不自动换行。"""
    for sid in ("SourceCode",):
        style = _find_style_by_id(doc, sid)
        if style is None:
            continue
        el = style.element
        pPr = el.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            el.insert(0, pPr)
        _set_spacing(pPr, before=0, after=0, line=260)
        _set_jc(pPr, "left")
        _clear_indent(pPr)
        # 关闭 wordWrap=off（pandoc 默认关闭导致行内不折行）
        ww = pPr.find(qn("w:wordWrap"))
        if ww is None:
            ww = OxmlElement("w:wordWrap")
            pPr.append(ww)
        ww.set(qn("w:val"), "on")
        rPr = _ensure_rPr(el)
        _set_rFonts(rPr, ascii="Consolas", east_asia="宋体")
        _set_sz(rPr, 18)
    # VerbatimChar 同步字体
    vc = _find_style_by_id(doc, "VerbatimChar")
    if vc is not None:
        rPr = _ensure_rPr(vc.element)
        _set_rFonts(rPr, ascii="Consolas")
        _set_sz(rPr, 18)


def override_list_styles(doc):
    """列表：匹配示例样式 ind=left=320 firstLine=400，并把有序列表改为中文圆圈数字。"""
    try:
        numbering_part = doc.part.numbering_part
    except Exception:
        return
    if numbering_part is None:
        return
    num_el = numbering_part.element
    # 所有 abstractNum 统一为示例间距 (hanging indent: left=320, firstLine=400)
    for ind in num_el.iter(qn("w:ind")):
        ind.set(qn("w:left"), "320")
        ind.set(qn("w:firstLine"), "400")
        # 清除 hanging/start 让只用 left+firstLine
        if ind.get(qn("w:hanging")) is not None:
            del ind.attrib[qn("w:hanging")]
    # 有序列表保持原生阿拉伯数字 "1."
    for lvl in num_el.iter(qn("w:lvl")):
        if lvl.get(qn("w:ilvl")) != "0":
            continue
        numFmt = lvl.find(qn("w:numFmt"))
        lvlText = lvl.find(qn("w:lvlText"))
        if numFmt is not None and lvlText is not None and numFmt.get(qn("w:val")) == "decimal":
            lvlText.set(qn("w:val"), "%1.")
        hanging = ind.get(qn("w:hanging"))
        if hanging and int(hanging) > 360:
            ind.set(qn("w:hanging"), "360")


def _iter_all_p(body):
    """遍历 body 下所有段落（包括 sdt/tbl 嵌套内）。"""
    for p in body.iter(qn("w:p")):
        yield p


def _p_text(p):
    return "".join(t.text or "" for t in p.iter(qn("w:t")))


def fix_abstract_layout(doc):
    """摘要/Abstract/目录 标题：黑体居中 14pt bold。
    关键词行：整段 bold，无首行缩进。"""
    body = doc.element.body
    for p in _iter_all_p(body):
        text = _p_text(p).strip()
        if text in ("摘要", "摘　要", "Abstract", "目录"):
            pPr = p.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                p.insert(0, pPr)
            # 摘要/Abstract 改为 Normal（避免被 TOC 收录），目录保持 TOCHeading
            if text in ("摘要", "摘　要", "Abstract"):
                # 把"摘　要"里的全角空格标准化成两个半角空格（匹配示例）
                if text == "摘　要":
                    for r in p.findall(qn("w:r")):
                        for t in r.findall(qn("w:t")):
                            if t.text and "　" in t.text:
                                t.text = t.text.replace("　", "  ")
                # 替换 pStyle 为 Normal
                ps = pPr.find(qn("w:pStyle"))
                if ps is None:
                    ps = OxmlElement("w:pStyle")
                    pPr.insert(0, ps)
                ps.set(qn("w:val"), "Normal")
                # 清除 outlineLvl（这是让 Heading 进 TOC 的关键属性）
                for ol in pPr.findall(qn("w:outlineLvl")):
                    pPr.remove(ol)
            _set_jc(pPr, "center")
            _clear_indent(pPr)
            for r in p.findall(qn("w:r")):
                rPr = _ensure_rPr(r)
                _set_rFonts(rPr, ascii="Times New Roman", east_asia="黑体")
                _set_sz(rPr, 28)
                _set_bold(rPr)
        elif text.startswith("关键词") or text.startswith("Keywords"):
            pPr = p.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                p.insert(0, pPr)
            _clear_indent(pPr)
            _set_jc(pPr, "left")
            for r in p.findall(qn("w:r")):
                rPr = _ensure_rPr(r)
                _set_bold(rPr)


def normalize_line_spacing(doc):
    """仅在段落 pPr 里出现 line=exact 且过大时改为 auto 360（保留 code block 等合理紧凑）。"""
    body = doc.element.body
    for p in _iter_all_p(body):
        pPr = p.find(qn("w:pPr"))
        if pPr is None:
            continue
        # 跳过代码块（SourceCode pStyle）
        ps = pPr.find(qn("w:pStyle"))
        if ps is not None and ps.get(qn("w:val")) == "SourceCode":
            continue
        sp = pPr.find(qn("w:spacing"))
        if sp is None:
            continue
        rule = sp.get(qn("w:lineRule"))
        line = sp.get(qn("w:line"))
        if rule == "exact" and line and int(line) > 420:
            sp.set(qn("w:line"), "360")
            sp.set(qn("w:lineRule"), "auto")


def apply_first_line_indent(doc):
    """段落级首行缩进控制：
    - 正文段落：强制加 2 字符首行缩进（覆盖 pandoc 的 ind=0）。
    - 列表项（带 numPr）/ 参考文献（[N] 开头）：显式清零首行缩进。
    - 其他特殊段落：跳过不动。
    """
    SKIP_STYLE_PREFIXES = ("Heading", "TOC", "Caption", "ImageCaption",
                           "CaptionedFigure", "SourceCode", "VerbatimChar",
                           "Title", "Subtitle")
    body = doc.element.body
    for p in _iter_all_p(body):
        pPr = p.find(qn("w:pPr"))
        style_name = ""
        if pPr is not None:
            ps = pPr.find(qn("w:pStyle"))
            if ps is not None:
                style_name = ps.get(qn("w:val")) or ""
        if any(style_name.startswith(pref) for pref in SKIP_STYLE_PREFIXES):
            continue
        # 跳过含图片/公式段落
        if p.find(qn("w:drawing")) is not None:
            continue
        text = _p_text(p).strip()
        if not text:
            continue
        if (text in ("摘要", "摘　要", "摘  要", "Abstract", "目录")
                or text.startswith("关键词") or text.startswith("Keywords")):
            continue
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p.insert(0, pPr)
        # 列表项和参考文献：首行缩进清零（覆盖 Normal/Compact 样式继承的缩进）
        is_list = pPr.find(qn("w:numPr")) is not None
        is_ref = bool(re.match(r"^\[\d+\]", text))
        if is_list or is_ref:
            _clear_indent(pPr)
        else:
            _set_indent(pPr, first_line_chars=200, first_line=480)


def insert_chapter_page_breaks(doc):
    """每个一级章节前插入 page break：摘要/Abstract/`N xxx`/参考文献。"""
    body = doc.element.body
    chapter_re = re.compile(r"^\d+\s+\S")
    SECTION_LABELS = {"摘要", "摘　要", "摘  要", "Abstract"}
    for p in list(_iter_all_p(body)):
        pPr = p.find(qn("w:pPr"))
        if pPr is None:
            continue
        ps = pPr.find(qn("w:pStyle"))
        style = ps.get(qn("w:val")) if ps is not None else ""
        text = _p_text(p).strip()
        # 摘要/Abstract 现在是 Normal+居中，通过文字匹配
        if text in SECTION_LABELS:
            pass  # 允许插入
        elif style == "Heading2" and (chapter_re.match(text) or text == "参考文献"):
            pass
        else:
            continue
        # 在段首插一个带 w:br type="page" 的 run
        first_r = p.find(qn("w:r"))
        if first_r is None:
            continue
        # 避免重复插入
        existing_br = first_r.find(qn("w:br"))
        if existing_br is not None and existing_br.get(qn("w:type")) == "page":
            continue
        new_r = OxmlElement("w:r")
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        new_r.append(br)
        p.insert(list(p).index(first_r), new_r)


def linkify_references(doc):
    """给参考文献每条加 bookmark `ref_N`，正文中的 [N] 改为 internal hyperlink。"""
    import re as _re
    body = doc.element.body
    bookmark_id = 5000
    ref_numbers = set()

    # —— Step 1: 参考文献章节的 [N] 加 bookmark ——
    in_refs = False
    for p in _iter_all_p(body):
        pPr = p.find(qn("w:pPr"))
        ps = pPr.find(qn("w:pStyle")) if pPr is not None else None
        style_val = ps.get(qn("w:val")) if ps is not None else ""
        text = _p_text(p)
        if style_val.startswith("Heading") and "参考文献" in text:
            in_refs = True
            continue
        if not in_refs:
            continue
        if style_val.startswith("Heading"):
            in_refs = False
            continue
        m = _re.match(r'\s*\[(\d+)\]', text)
        if not m:
            continue
        n = m.group(1)
        if n in ref_numbers:
            continue
        ref_numbers.add(n)
        bms = OxmlElement("w:bookmarkStart")
        bms.set(qn("w:id"), str(bookmark_id))
        bms.set(qn("w:name"), f"ref_{n}")
        bme = OxmlElement("w:bookmarkEnd")
        bme.set(qn("w:id"), str(bookmark_id))
        bookmark_id += 1
        first_r = p.find(qn("w:r"))
        if first_r is not None:
            idx = list(p).index(first_r)
            p.insert(idx, bms)
            p.insert(idx + 1, bme)

    # —— Step 2: 正文里的 [N] 替换为 hyperlink ——
    pat = _re.compile(r'\[(\d+)\]')
    in_refs = False
    for p in _iter_all_p(body):
        pPr = p.find(qn("w:pPr"))
        ps = pPr.find(qn("w:pStyle")) if pPr is not None else None
        style_val = ps.get(qn("w:val")) if ps is not None else ""
        text = _p_text(p)
        if style_val.startswith("Heading") and "参考文献" in text:
            in_refs = True
            continue
        if in_refs:
            # 不在参考文献段落中转换 [N]，避免把编号本身变成自己的链接
            continue
        # 跳过代码块
        if style_val == "SourceCode":
            continue
        # 迭代 run，替换 [N]
        changed = True
        while changed:
            changed = False
            for r in list(p.findall(qn("w:r"))):
                ts = r.findall(qn("w:t"))
                if len(ts) != 1:
                    continue
                t = ts[0]
                txt = t.text or ""
                m = pat.search(txt)
                if not m:
                    continue
                n = m.group(1)
                if n not in ref_numbers:
                    continue  # 只链接实际存在的参考文献
                before, after = txt[:m.start()], txt[m.end():]
                rPr_src = r.find(qn("w:rPr"))
                parent = r.getparent()
                idx = list(parent).index(r)
                parent.remove(r)

                def _clone_run(text_content):
                    new_r = OxmlElement("w:r")
                    if rPr_src is not None:
                        new_r.append(copy.deepcopy(rPr_src))
                    new_t = OxmlElement("w:t")
                    new_t.set(qn("xml:space"), "preserve")
                    new_t.text = text_content
                    new_r.append(new_t)
                    return new_r

                # 国标引用格式：黑色上标（保留点击跳转）
                hp = OxmlElement("w:hyperlink")
                hp.set(qn("w:anchor"), f"ref_{n}")
                hp.set(qn("w:history"), "1")
                link_r = OxmlElement("w:r")
                if rPr_src is not None:
                    link_rPr = copy.deepcopy(rPr_src)
                else:
                    link_rPr = OxmlElement("w:rPr")
                for c in link_rPr.findall(qn("w:color")):
                    link_rPr.remove(c)
                for u in link_rPr.findall(qn("w:u")):
                    link_rPr.remove(u)
                for v in link_rPr.findall(qn("w:vertAlign")):
                    link_rPr.remove(v)
                col = OxmlElement("w:color")
                col.set(qn("w:val"), "auto")
                link_rPr.append(col)
                va = OxmlElement("w:vertAlign")
                va.set(qn("w:val"), "superscript")
                link_rPr.append(va)
                link_r.append(link_rPr)
                link_t = OxmlElement("w:t")
                link_t.set(qn("xml:space"), "preserve")
                link_t.text = f"[{n}]"
                link_r.append(link_t)
                hp.append(link_r)

                insert_at = idx
                if after:
                    parent.insert(insert_at, _clone_run(after))
                parent.insert(insert_at, hp)
                if before:
                    parent.insert(insert_at, _clone_run(before))
                changed = True
                break  # restart scan


def enable_update_fields_on_open(doc):
    """让 Word/LibreOffice 首次打开时自动更新 TOC 等域。"""
    settings = doc.settings.element
    # 移除已有的 updateFields
    for old in settings.findall(qn("w:updateFields")):
        settings.remove(old)
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    settings.insert(0, upd)


def fix_body_indent(doc):
    """给正文段落加 2 字符首行缩进（不影响标题/图/列表/TOC）。"""
    skip_styles = {"Heading1", "Heading2", "Heading3", "Heading4", "Heading5",
                   "Heading 1", "Heading 2", "Heading 3", "Heading 4",
                   "CaptionedFigure", "ImageCaption", "Caption",
                   "TOC1", "TOC2", "TOC3", "TOCHeading",
                   "TOC 1", "TOC 2", "TOC 3", "TOC Heading",
                   "SourceCode", "VerbatimChar", "CompactFirstLineIndent"}
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ""
        if style_name in skip_styles or style_name.startswith(("Heading", "TOC", "Caption")):
            continue
        text = p.text.strip()
        if not text:
            continue
        if text.startswith("关键词") or text.startswith("Keywords"):
            continue
        # 跳过含图/公式的段落（通常无 text 或 text 只含 caption）
        if p._element.find(qn("w:drawing")) is not None:
            continue
        pPr = _ensure_pPr(p)
        # 只有没设过 ind 时才加
        if pPr.find(qn("w:ind")) is None:
            _set_indent(pPr, first_line_chars=200, first_line=480)


def main():
    # Step 1: pandoc 生成主体
    pandoc_out = REPORT_DIR / "_pandoc_out.docx"
    run_pandoc(MD_FILE, pandoc_out)
    print(f"Pandoc →  {pandoc_out}")

    # Step 2: python-docx 后处理
    doc = Document(str(pandoc_out))

    setup_page(doc)
    setup_header_footer(doc)
    override_heading_styles(doc)
    override_body_styles(doc)
    override_code_style(doc)
    override_list_styles(doc)
    normalize_line_spacing(doc)
    fix_abstract_layout(doc)
    apply_first_line_indent(doc)
    insert_chapter_page_breaks(doc)
    linkify_references(doc)
    enable_update_fields_on_open(doc)

    # Step 3: 封面
    insert_cover(doc, TEMPLATE_FILE)

    doc.save(str(OUTPUT_FILE))
    pandoc_out.unlink(missing_ok=True)
    print(f"Saved → {OUTPUT_FILE}")
    print(f"Size: {OUTPUT_FILE.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    main()

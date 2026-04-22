# /// script
# requires-python = ">=3.12"
# dependencies = ["python-docx>=1.1"]
# ///
"""原地修补 report.docx：只替换指定字符串，不重建文档，保留手工编辑与格式。

用法：uv run patch_docx.py
"""
from pathlib import Path
from docx import Document

DOCX = Path(__file__).parent / "report.docx"

# (旧字符串, 新字符串) — 按出现顺序排列
PATCHES = [
    # Soundex 运行界面段落
    (
        "flo → F400 → fail, fale, fall, feel, fell",
        "flo → F400 → flow, fli, fl, fail, fale",
    ),
    (
        "（发音相近但不全相关，体现了 Soundex 召回为主的特点）",
        "（候选按与原词公共前缀长度排序，最相关的 flow 排在首位）",
    ),
    (
        "再用扩展后的词集参与 TF-IDF 排序得到 50 条排名靠前的检索结果。",
        "检索阶段只取每个查询词的 top-1 候选参与 TF-IDF 排序，"
        "避免拼写相近但语义无关的候选污染结果；最终得到 50 条排序后的文档。",
    ),
    # suggest 代码块：旧一行改为按公共前缀排序的一行
    (
        'candidates = sorted(self.code_to_terms.get(code, set()) - {word.lower()})',
        'raw = self.code_to_terms.get(code, set()) - {word.lower()}\n'
        '        candidates = sorted(raw, key=lambda c: (-_common_prefix(word, c), abs(len(c)-len(word)), c))',
    ),
]


def patch_runs(doc: Document) -> int:
    """遍历所有段落的 run 做 text 替换，返回替换次数。"""
    n = 0
    # 正文段落
    for p in doc.paragraphs:
        n += _patch_para(p)
    # 表格里的段落
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    n += _patch_para(p)
    return n


def _patch_para(p) -> int:
    """对单段所有 run 做替换；若跨 run 则合并后替换。"""
    n = 0
    # 先尝试每个 run 独立替换（保留格式）
    for run in p.runs:
        for old, new in PATCHES:
            if old in run.text:
                run.text = run.text.replace(old, new)
                n += 1
    # 再尝试整段 text 匹配（跨 run）— 将匹配到的 PATCH 整体合并到第一个 run
    full = p.text
    for old, new in PATCHES:
        if old in full and all(old not in r.text for r in p.runs):
            # 跨 run：把整段文字合到首个 run，清空其余
            if p.runs:
                first = p.runs[0]
                first.text = full.replace(old, new)
                for r in p.runs[1:]:
                    r.text = ""
                n += 1
                break
    return n


def main():
    doc = Document(str(DOCX))
    n = patch_runs(doc)
    print(f"patches applied: {n}/{len(PATCHES)}")
    if n == 0:
        print("  (no matches — patches may have already been applied or text differs)")
        return
    doc.save(str(DOCX))
    print(f"saved → {DOCX}")


if __name__ == "__main__":
    main()
